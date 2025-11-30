"""
리포트 생성 모듈
python-docx를 사용하여 DOCX 형식의 기술본 리포트를 생성합니다.
"""

import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

from src.config import PROJECT_ROOT
from src.llm import LLMReportGenerator

logger = logging.getLogger(__name__)


class ReportGenerator:
    """리포트 생성 클래스"""

    def __init__(self, llm_generator: Optional[LLMReportGenerator] = None):
        """
        Args:
            llm_generator: LLM 리포트 생성기 (None이면 자동 생성)
        """
        # LLM 생성기 초기화 (명시적으로 생성하여 로그 확인)
        if llm_generator is None:
            import logging
            logger = logging.getLogger(__name__)
            logger.info("LLMReportGenerator 초기화 중...")
            llm_generator = LLMReportGenerator()
            if llm_generator.client:
                logger.info("✅ LLM 연결 성공")
            else:
                logger.warning("⚠️ LLM 연결 실패 - Fallback summary 사용")
        self.llm_generator = llm_generator
        self.reports_dir = PROJECT_ROOT / "reports"
        self.reports_dir.mkdir(parents=True, exist_ok=True)

    def generate_report(
        self,
        report_id: str,
        scan_results: List[Dict[str, Any]],
        poc_reproductions: List[Dict[str, Any]],
        executive_summary: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        리포트 생성

        Args:
            report_id: 리포트 ID
            scan_results: 스캔 결과 리스트
            poc_reproductions: PoC 재현 결과 리스트
            executive_summary: Executive Summary (None이면 LLM으로 생성)

        Returns:
            리포트 생성 결과 딕셔너리
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # DOCX 문서 생성
            doc = Document()

            # 1. 표지
            self._add_cover_page(doc, report_id)

            # 2. Executive Summary
            if executive_summary is None:
                try:
                    summary_data = self.llm_generator.generate_executive_summary(
                        scan_results, poc_reproductions
                    )
                    executive_summary = summary_data.get("executive_summary", "")
                    if not executive_summary:
                        logger.warning("LLM이 Executive Summary를 생성하지 못했습니다. Fallback summary를 사용합니다.")
                        executive_summary = "LLM을 통한 Executive Summary 생성에 실패했습니다. 기본 요약을 사용합니다."
                except Exception as e:
                    logger.error(f"Executive Summary 생성 중 오류: {str(e)}")
                    executive_summary = f"Executive Summary 생성 중 오류가 발생했습니다: {str(e)}"

            self._add_executive_summary(doc, executive_summary)

            # 3. 취약점 상세
            self._add_vulnerability_details(
                doc, scan_results, poc_reproductions
            )

            # 4. 증거 참조
            self._add_evidence_references(doc, poc_reproductions)

            # 5. 권고사항
            self._add_recommendations(doc, scan_results)

            # 리포트 저장
            report_path = self.reports_dir / f"{report_id}.docx"
            doc.save(str(report_path))

            logger.info(f"Report generated: {report_path}")

            return {
                "success": True,
                "report_id": report_id,
                "file_path": str(report_path),
                "file_size": report_path.stat().st_size
            }

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {
                "success": False,
                "error": "python-docx package not installed"
            }
        except Exception as e:
            logger.error(f"Failed to generate report: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }

    def _add_cover_page(self, doc, report_id: str):
        """표지 페이지 추가"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title = doc.add_heading("취약점 진단 보고서", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(f"Report ID: {report_id}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_executive_summary(self, doc, summary: str):
        """Executive Summary 추가"""
        doc.add_heading("Executive Summary", 1)

        # 요약 내용 추가
        for line in summary.split("\n"):
            if line.strip():
                if line.startswith("#"):
                    # 마크다운 헤딩을 DOCX 헤딩으로 변환
                    level = line.count("#")
                    text = line.lstrip("#").strip()
                    doc.add_heading(text, level=min(level, 3))
                else:
                    doc.add_paragraph(line.strip())

        doc.add_page_break()

    def _add_vulnerability_details(
        self,
        doc,
        scan_results: List[Dict[str, Any]],
        poc_reproductions: List[Dict[str, Any]]
    ):
        """취약점 상세 추가"""
        doc.add_heading("취약점 상세", 1)

        for scan_result in scan_results:
            normalized = scan_result.get("normalized_result", {})
            findings = normalized.get("findings", [])

            for finding in findings:
                # 취약점 제목
                doc.add_heading(finding.get("title", "Unknown"), 2)

                # 기본 정보
                doc.add_paragraph(f"심각도: {finding.get('severity', 'Unknown')}")
                doc.add_paragraph(f"CVE: {', '.join(finding.get('cve_list', []))}")
                doc.add_paragraph(f"설명: {finding.get('description', '')}")

                # PoC 재현 결과 연결
                related_pocs = [
                    p for p in poc_reproductions
                    if p.get("scan_result_id") == scan_result.get("id")
                ]

                if related_pocs:
                    doc.add_paragraph("PoC 재현 결과:")
                    for poc in related_pocs:
                        doc.add_paragraph(
                            f"- 상태: {poc.get('status', 'Unknown')}, "
                            f"신뢰도: {poc.get('reliability_score', 'N/A')}",
                            style="List Bullet"
                        )

                doc.add_paragraph()  # 빈 줄

    def _add_evidence_references(self, doc, poc_reproductions: List[Dict[str, Any]]):
        """증거 참조 추가"""
        doc.add_heading("증거 참조", 1)

        for poc in poc_reproductions:
            doc.add_paragraph(f"PoC ID: {poc.get('reproduction_id', 'Unknown')}")

            if poc.get("syscall_log_path"):
                doc.add_paragraph(f"- 시스템콜 로그: {poc['syscall_log_path']}")

            if poc.get("network_capture_path"):
                doc.add_paragraph(f"- 네트워크 캡처: {poc['network_capture_path']}")

            if poc.get("filesystem_diff_path"):
                doc.add_paragraph(f"- 파일 시스템 변화: {poc['filesystem_diff_path']}")

            doc.add_paragraph()

    def _add_recommendations(self, doc, scan_results: List[Dict[str, Any]]):
        """권고사항 추가"""
        doc.add_heading("권고사항", 1)

        # 심각도별 권고사항
        doc.add_heading("Critical/High 심각도 취약점", 2)
        doc.add_paragraph("1. 즉시 패치 적용")
        doc.add_paragraph("2. 임시 조치 적용 (방화벽 규칙, 접근 제한 등)")
        doc.add_paragraph("3. 모니터링 강화")

        doc.add_heading("Medium/Low 심각도 취약점", 2)
        doc.add_paragraph("1. 계획된 패치 일정 수립")
        doc.add_paragraph("2. 정기적인 보안 점검")
        doc.add_paragraph("3. 보안 설정 최적화")

